#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Word文档处理工具
基于word_template_processor.py，为SuitAgent系统封装

版本: v1.0
更新: 2025-11-14
作者: SuitAgent整合系统
"""

import os
import sys
import logging
from typing import Dict, List, Tuple, Optional
from pathlib import Path

# 设置日志
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

try:
    from docx import Document
except ImportError:
    logger.error("python-docx 未安装，请运行: pip install python-docx")
    sys.exit(1)


class Execute:
    """
    段落占位符替换执行器
    基于原系统的Execute类，为SuitAgent优化

    该类负责在不破坏格式的前提下，替换Word文档中的占位符
    核心技术：run级别操作，保持原始格式
    """

    def __init__(self, paragraph):
        """初始化执行器

        Args:
            paragraph: docx段落对象
        """
        self.paragraph = paragraph

    def p_replace(self, x: int, key: str, value: str):
        """
        执行段落替换（保持格式）

        重要：不直接替换段落文本，而是操作run对象，以保持原始格式

        Args:
            x: 段落索引（用于日志）
            key: 占位符名称（包含大括号，如 {client}）
            value: 替换值

        Raises:
            ValueError: 当占位符或值为空时
        """
        if not key or not value:
            logger.warning(f"占位符或值为空，跳过: {key}")
            return

        try:
            # 获取段落中所有字符的坐标索引
            # 格式: [{"run": run_index, "char": char_index}, ...]
            p_maps = [
                {"run": y, "char": z}
                for y, run in enumerate(self.paragraph.runs)
                for z, char in enumerate(list(run.text))
            ]

            # 查找占位符在段落中的所有起始位置
            k_idx = [
                s for s in range(len(self.paragraph.text))
                if self.paragraph.text.find(key, s, len(self.paragraph.text)) == s
            ]

            if not k_idx:
                logger.debug(f"段落 {x}: 未找到占位符 {key}")
                return

            # 逆序替换，避免索引偏移问题
            for i, start_idx in enumerate(reversed(k_idx)):
                end_idx = start_idx + len(key)
                k_maps = p_maps[start_idx:end_idx]

                if k_maps:
                    self.r_replace(k_maps, value)

            logger.debug(f"段落 {x}: 成功替换 {key} => {value}")

        except Exception as e:
            logger.error(f"段落 {x}: 替换失败 {key} => {value}, 错误: {e}")

    def r_replace(self, k_maps: List[dict], value: str):
        """
        运行级替换（保持格式）

        通过反向迭代删除和替换字符，避免索引错误

        Args:
            k_maps: 占位符的坐标映射列表
                    例如: [{"run": 15, "char": 3}, {"run": 15, "char": 4}]
            value: 替换值
        """
        try:
            # 逆序迭代，处理每个字符
            for i, position in enumerate(reversed(k_maps), start=1):
                run_index = position["run"]
                char_index = position["char"]

                # 获取run对象
                if run_index >= len(self.paragraph.runs):
                    logger.warning(f"run索引超出范围: {run_index}")
                    continue

                run = self.paragraph.runs[run_index]

                # 将run文本转换为字符列表
                # 避免某些情况下run.text出现错误
                thisrun = list(run.text)

                if i < len(k_maps):
                    # 非最后一个字符，删除
                    if char_index < len(thisrun):
                        thisrun.pop(char_index)
                else:
                    # 最后一个字符，替换
                    if char_index < len(thisrun):
                        thisrun[char_index] = value

                # 重新设置run文本
                run.text = ''.join(thisrun)

        except Exception as e:
            logger.error(f"run替换失败: {e}")


class DocxProcessor:
    """
    Word文档处理器

    提供完整的Word文档占位符替换功能，支持段落和表格
    """

    @staticmethod
    def body_content(doc, replace_dict: Dict[str, str]):
        """
        处理文档正文中的占位符

        Args:
            doc: docx文档对象
            replace_dict: 占位符字典 {占位符: 值}
        """
        logger.info("  ☺ 处理文档正文中的占位符...")
        replaced_count = 0

        for key, value in replace_dict.items():
            placeholder = f"{{{key}}}"
            for x, paragraph in enumerate(doc.paragraphs):
                if placeholder in paragraph.text:
                    Execute(paragraph).p_replace(x, placeholder, str(value))
                    replaced_count += 1

        logger.info(f"  ✓ 正文占位符处理完成，共处理 {replaced_count} 处")

    @staticmethod
    def body_tables(doc, replace_dict: Dict[str, str]):
        """
        处理文档表格中的占位符

        Args:
            doc: docx文档对象
            replace_dict: 占位符字典 {占位符: 值}
        """
        logger.info("  ☺ 处理文档表格中的占位符...")
        replaced_count = 0

        for key, value in replace_dict.items():
            placeholder = f"{{{key}}}"
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for x, paragraph in enumerate(cell.paragraphs):
                            if placeholder in paragraph.text:
                                Execute(paragraph).p_replace(x, placeholder, str(value))
                                replaced_count += 1

        logger.info(f"  ✓ 表格占位符处理完成，共处理 {replaced_count} 处")

    @staticmethod
    def process_document(
        template_path: str,
        output_path: str,
        replace_dict: Dict[str, str]
    ) -> bool:
        """
        处理单个Word文档模板，替换占位符并生成新文档

        Args:
            template_path: 模板文件路径
            output_path: 输出文件路径
            replace_dict: 占位符字典 {占位符: 值}

        Returns:
            bool: 是否成功

        Raises:
            FileNotFoundError: 当模板文件不存在时
            Exception: 当处理过程中发生错误时
        """
        try:
            # 验证模板文件
            if not os.path.exists(template_path):
                raise FileNotFoundError(f"模板文件不存在: {template_path}")

            logger.info(f"开始处理文档: {template_path}")

            # 打开模板文件
            doc = Document(template_path)

            # 处理文档内容和表格中的占位符
            DocxProcessor.body_content(doc, replace_dict)
            DocxProcessor.body_tables(doc, replace_dict)

            # 确保输出目录存在
            output_dir = os.path.dirname(output_path)
            if output_dir and not os.path.exists(output_dir):
                os.makedirs(output_dir)
                logger.info(f"创建输出目录: {output_dir}")

            # 保存生成的文档
            doc.save(output_path)
            logger.info(f"✓ 文档生成成功: {output_path}")

            return True

        except Exception as e:
            logger.error(f"✗ 文档处理失败: {e}", exc_info=True)
            return False

    @staticmethod
    def process_from_yaml(
        template_path: str,
        output_path: str,
        yaml_file_path: str
    ) -> bool:
        """
        从YAML文件生成Word文档

        Args:
            template_path: Word模板路径
            output_path: 输出文件路径
            yaml_file_path: 案件YAML文件路径

        Returns:
            bool: 是否成功
        """
        try:
            # 导入placeholder_mapper
            from placeholder_mapper import PlaceholderMapper

            # 加载YAML文件
            yaml_data = PlaceholderMapper.load_yaml_from_file(yaml_file_path)

            # 转换为占位符
            placeholders = PlaceholderMapper.yaml_to_placeholders(yaml_data)

            # 验证必填字段
            missing = PlaceholderMapper.validate_required_fields(placeholders)
            if missing:
                logger.warning(f"缺失必填字段: {', '.join(missing)}")

            # 处理文档
            return DocxProcessor.process_document(
                template_path, output_path, placeholders
            )

        except ImportError:
            logger.error("placeholder_mapper 未找到，请确保文件存在")
            return False
        except Exception as e:
            logger.error(f"从YAML处理失败: {e}", exc_info=True)
            return False

    @staticmethod
    def process_directory(
        template_dir: str,
        output_dir: str,
        replace_dict: Dict[str, str]
    ) -> List[Tuple[str, bool]]:
        """
        批量处理目录中的所有Word文档模板

        Args:
            template_dir: 模板目录路径
            output_dir: 输出目录路径
            replace_dict: 占位符字典

        Returns:
            List[Tuple[str, bool]]: 处理结果列表 [(文件名, 是否成功), ...]
        """
        results = []

        try:
            if not os.path.exists(template_dir):
                raise FileNotFoundError(f"模板目录不存在: {template_dir}")

            if not os.path.exists(output_dir):
                os.makedirs(output_dir)

            # 获取所有docx文件
            files = os.listdir(template_dir)
            doc_files = [f for f in files if f.endswith(('.docx', '.doc'))]

            logger.info(f"找到 {len(doc_files)} 个模板文件")

            for file in doc_files:
                template_path = os.path.join(template_dir, file)
                output_path = os.path.join(output_dir, file)

                success = DocxProcessor.process_document(
                    template_path, output_path, replace_dict
                )
                results.append((file, success))

            # 统计结果
            success_count = sum(1 for _, success in results if success)
            logger.info(f"✓ 批量处理完成: {success_count}/{len(doc_files)} 成功")

            return results

        except Exception as e:
            logger.error(f"批量处理失败: {e}", exc_info=True)
            return results


def main():
    """主函数 - 用于测试"""
    print("=" * 80)
    print("DocxProcessor 测试")
    print("=" * 80)

    # 创建测试数据
    test_placeholders = {
        'client': '张三有限公司',
        'type': '公司',
        'lawyer': '王五律师',
        'date': '2025-11-14'
    }

    print("\n测试占位符:")
    for key, value in test_placeholders.items():
        print(f"  {{{key}}}: {value}")

    print("\n注意: 需要准备实际的Word模板文件才能完成测试")
    print("=" * 80)


if __name__ == "__main__":
    main()
